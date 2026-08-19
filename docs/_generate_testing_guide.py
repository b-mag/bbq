"""Generate docs/Developer_Testing_Guide.docx for new clones."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Developer_Testing_Guide.docx"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=18 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    return p


def screenshot_box(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(f"[SCREENSHOT PLACEHOLDER] {label}")
    set_run_font(run, size=10, italic=True, color=RGBColor(0x88, 0x44, 0x00))
    p.paragraph_format.space_after = Pt(12)
    return p


def script_section(doc, title, bat_path, when, launches, expect, status):
    add_heading(doc, title, 2)
    add_para(doc, f"File: {bat_path}", bold=True)
    add_para(doc, "When to use it")
    add_bullet(doc, when)
    add_para(doc, "What it launches")
    for item in launches:
        add_bullet(doc, item)
    add_para(doc, "What you should see")
    for item in expect:
        add_bullet(doc, item)
    add_para(doc, f"Status: {status}", bold=True)
    screenshot_box(doc, title)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Carcosa Developer Testing Guide")
    set_run_font(run, size=26, bold=True)
    sub = doc.add_paragraph()
    run = sub.add_run("Get a clone running, then verify peer mesh scripts. Screenshot slots are left for confirmation.")
    set_run_font(run, size=12, italic=True, color=RGBColor(0x44, 0x44, 0x44))

    add_heading(doc, "1. Who this is for", 1)
    add_para(
        doc,
        "This guide is for a developer who just cloned the repo. It is not a player distribution readme. "
        "It tells you which build to run first, which double-click scripts to use, and what each combination is supposed to do.",
    )
    add_bullet(doc, "Release scripts (Native AOT publish) are what most testing uses today.")
    add_bullet(doc, "Develop scripts do the same launches against Debug builds.")
    add_bullet(doc, "Python scripts under dev_scripts/python/ pack art. They do not launch the game.")

    add_heading(doc, "2. Prerequisites", 1)
    add_bullet(doc, "Windows 10/11 x64")
    add_bullet(doc, ".NET 10 SDK")
    add_bullet(doc, "Node.js 18+ and npm (frontends are built into the game and matchmaking exes)")
    add_bullet(doc, "Optional: Docker, only if you want Kafka + compose matchmaking instead of the local tracker exe")

    add_heading(doc, "3. First step: build the release binaries", 1)
    add_para(
        doc,
        "Run build_all_release.bat first. This produces Native AOT release builds that most of the scripts below use. "
        "It takes several minutes. Do this after a clone, and again after gameplay or UI changes you want to test in the published exes.",
        bold=True,
    )
    add_code(doc, r"dev_scripts\release\build_all_release.bat")
    add_para(doc, "What it builds")
    add_bullet(doc, r"src\backend\bin\Release\net10.0\win-x64\publish\Carcosa.exe  (game + embedded Next.js UI)")
    add_bullet(doc, r"src\matchmaking\bin\Release\net10.0\win-x64\publish\Carcosa.Matchmaking.exe  (optional tracker + dashboard)")
    add_bullet(doc, r"src\botclient\bin\Release\net10.0\win-x64\publish\Carcosa.BotClient.exe")
    add_para(
        doc,
        "If a launch script says the exe was not found, you skipped this step or the publish failed. "
        "Do not start from launch-full-test.bat on a fresh clone.",
    )
    screenshot_box(doc, "build_all_release.bat console: ALL BUILDS SUCCEEDED")

    add_heading(doc, "4. Folder map", 1)
    add_bullet(doc, r"dev_scripts\release\  - AOT publish + launchers used for the testing described here")
    add_bullet(doc, r"dev_scripts\develop\  - Debug equivalents. Run build_all_debug.bat first, then the same launcher names.")
    add_bullet(doc, r"dev_scripts\python\   - palettes / sprites / tilesets art pipeline")
    add_bullet(doc, r"dev_scripts\obsolete_or_fix\  - old per-project builders; ignore")
    add_bullet(doc, r"docs\  - current technical notes and this guide")
    add_bullet(doc, r"implementations\  - plans and future-work writeups")
    add_bullet(doc, r"backlog\  - ticket-sized stories")

    add_heading(doc, "5. How local peers find each other", 1)
    add_para(doc, "There are three discovery paths. Scripts pick a combination on purpose.")
    add_bullet(doc, "Tracker (matchmaking on port 5100): peers register and get each other's addresses. Fastest local auto-connect.")
    add_bullet(doc, "Glyph sharing: player copies a Glyph from the UI and the other player pastes it. Required when the tracker is off.")
    add_bullet(doc, "Peer cache (known-peers.json): remembers previous IPs. Useful after a real internet session; dangerous on localhost because leftover WAN IPs fail NAT hairpin.")
    add_para(
        doc,
        "Same-machine tests pin 127.0.0.1. Production / long-distance Glyph tests must launch Carcosa.exe WITHOUT --public-address so STUN can advertise the real public IP.",
    )

    add_heading(doc, "6. Release launch scripts", 1)
    add_para(
        doc,
        "Double-click the .bat in Windows Explorer. A console window stays open. Press any key in that window to stop everything the script started.",
    )

    script_section(
        doc,
        "launch-full-test.bat",
        r"dev_scripts\release\launch-full-test.bat",
        "Smoke the whole local stack the way a published build runs: tracker plus two peers, no localhost pin. Closest to 'just run the exes'.",
        [
            "Carcosa.Matchmaking.exe on http://localhost:5100",
            "Carcosa.exe on http://localhost:5000 with one bot",
            "Carcosa.exe on http://localhost:5001",
        ],
        [
            "Three windows open (tracker dashboard + two game UIs).",
            "Peers may try STUN/public IPs. On one PC that can fail NAT hairpin. If they do not see each other, use a localhost-pinned tracker script instead.",
            "Player 1 has a spawned bot.",
        ],
        "Local testing confirmed (launcher opens the stack). Mesh join on the same PC may still need a pinned-localhost script.",
    )

    script_section(
        doc,
        "launch-two-players-local-tracker.bat",
        r"dev_scripts\release\launch-two-players-local-tracker.bat",
        "Test optional tracker discovery on one machine. Peers should auto-connect without pasting a Glyph.",
        [
            "Tracker on port 5100",
            "Franz on port 5000 with --public-address=127.0.0.1:5000",
            "Marina on port 5001 with --public-address=127.0.0.1:5001",
            "Both pointed at http://127.0.0.1:5100",
        ],
        [
            "Both UIs open.",
            "Console logs show pinned 127.0.0.1, not your WAN IP.",
            "Within a few seconds the tracker registers both peers and they appear in each other's world.",
            "Peer cache is still used. Stale WAN IPs from an internet test can interfere; use the no-cache variant if that happens.",
        ],
        "Local testing confirmed.",
    )

    script_section(
        doc,
        "launch-two-players-local-tracker-no-cache.bat",
        r"dev_scripts\release\launch-two-players-local-tracker-no-cache.bat",
        "Same as the local tracker test, but do not reuse previous connections. Use after a long-distance Glyph test so known-peers.json cannot dial old WAN IPs.",
        [
            "Same tracker + Franz + Marina as the local-tracker script",
            "--no-cache-connect and --clear-peer-cache on both peers",
        ],
        [
            "Tracker still auto-discovers both peers.",
            "known-peers.json is cleared; they should not connect via leftover IPs.",
            "Logs still show 127.0.0.1.",
        ],
        "Pending confirmation (awaiting screenshots).",
    )

    script_section(
        doc,
        "launch-two-players-no-tracker-no-cache.bat",
        r"dev_scripts\release\launch-two-players-no-tracker-no-cache.bat",
        "Run two peer instances that do not auto-connect from previous connections. No tracker is started, so they must connect via Glyph sharing in game.",
        [
            "Franz on port 5000, Marina on port 5001",
            "Tracker URL forced to a dead address (http://127.0.0.1:1)",
            "Loopback Glyphs (--public-address=127.0.0.1:port)",
            "--no-cache-connect and --clear-peer-cache",
        ],
        [
            "Two game windows. They will NOT find each other automatically.",
            "In Player 1, copy the Glyph. In Player 2, join using that Glyph (or the reverse).",
            "After a successful join both players should appear on the overworld.",
            "Confirm logs show 127.0.0.1, not the WAN IP.",
        ],
        "Local testing confirmed.",
    )

    script_section(
        doc,
        "launch-two-players.bat",
        r"dev_scripts\release\launch-two-players.bat",
        "Glyph / manual discovery with tracker disabled, but peer cache still allowed. Use to test Glyph plus cache bootstrap.",
        [
            "Franz :5000 and Marina :5001",
            "Dead tracker URL",
            "Loopback public addresses",
            "Cache bootstrap left on",
        ],
        [
            "They should not auto-join through matchmaking.",
            "Join with Glyphs, or they may reconnect if known-peers.json already has 127.0.0.1 entries.",
        ],
        "Pending confirmation (awaiting screenshots).",
    )

    script_section(
        doc,
        "launch-two-players-no-cache.bat",
        r"dev_scripts\release\launch-two-players-no-cache.bat",
        "Older sibling of no-tracker-no-cache. Tracker off, cache cleared, Glyph required. Prefer launch-two-players-no-tracker-no-cache.bat for new tests; this one is kept for the same behavior with the original name.",
        [
            "Two localhost peers, dead tracker, --no-cache-connect, --clear-peer-cache",
        ],
        [
            "Same expectation as no-tracker-no-cache: Glyph sharing required.",
        ],
        "Pending confirmation (awaiting screenshots).",
    )

    script_section(
        doc,
        "prepare-release.bat",
        r"dev_scripts\release\prepare-release.bat",
        "Copy the published Carcosa.exe into a clean carcosa-release/ folder with appsettings for handing a build to a remote tester. Run build_all_release.bat first.",
        [
            "Does not launch the game.",
            "Writes Carcosa.exe, wwwroot, appsettings.json, and an offline config next to it.",
        ],
        [
            "Folder carcosa-release/ appears at the repo root (unless you pass another output path).",
            "Remote testers double-click Carcosa.exe. For internet Glyphs, do not pass --public-address.",
        ],
        "Pending confirmation (awaiting screenshots).",
    )

    add_heading(doc, "7. Develop (Debug) counterparts", 1)
    add_para(
        doc,
        "Same launcher names live in dev_scripts\\develop\\. They point at Debug exes instead of the AOT publish folder:",
    )
    add_code(doc, r"src\backend\bin\Debug\net10.0\Carcosa.exe")
    add_code(doc, r"src\matchmaking\bin\Debug\net10.0\Carcosa.Matchmaking.exe")
    add_para(doc, "First step for that folder:")
    add_code(doc, r"dev_scripts\develop\build_all_debug.bat")
    add_para(
        doc,
        "Debug builds include the React frontends but skip Native AOT, so they iterate faster. Behavior of each launch-*.bat should match the release twin. Status of the develop copies: pending confirmation.",
    )
    screenshot_box(doc, "build_all_debug.bat console: ALL BUILDS SUCCEEDED (Debug)")

    add_heading(doc, "8. Glyph sharing walkthrough", 1)
    add_bullet(doc, "Start two peers with a no-tracker script (or any script that does not auto-join).")
    add_bullet(doc, "In the UI, open the Glyph / join control and copy Player 1's Glyph.")
    add_bullet(doc, "On Player 2, paste and connect.")
    add_bullet(doc, "Local loopback Glyphs decode to 127.0.0.1:5000 or :5001. Internet Glyphs encode the STUN-mapped public address.")
    add_bullet(doc, "Long-distance success has been confirmed by sharing a Glyph between machines without --public-address.")
    screenshot_box(doc, "Player 1 Glyph copied; Player 2 joined; both visible on overworld")

    add_heading(doc, "9. Internet / remote testing", 1)
    add_para(
        doc,
        "Production Carcosa.exe discovers a public address with STUN, optionally registers with a tracker, and falls back to Glyphs. "
        "Do not pin 127.0.0.1 for a remote tester.",
    )
    add_bullet(doc, "Host: publish with build_all_release.bat, optionally prepare-release.bat, run Carcosa.exe with no --public-address.")
    add_bullet(doc, "Optional tracker: run Carcosa.Matchmaking.exe on a reachable host and point peers with --matchmaking-url=http://HOST:5100")
    add_bullet(doc, "Firewall: allow TCP on the listen port (default 5000) inbound. UPnP is attempted automatically.")
    add_bullet(doc, "If two home NATs are both symmetric / CGNAT, direct connect can still fail. TURN relay is not implemented yet (see implementations/NAT_TURN_GAP.md).")
    add_bullet(doc, "Offline mode in Settings skips the tracker; Glyphs still work.")

    add_heading(doc, "10. Command-line flags the scripts use", 1)
    add_bullet(doc, "--port=N                 HTTP / WebSocket listen port (5000, 5001)")
    add_bullet(doc, "--name=X                 Display name (Franz / Marina)")
    add_bullet(doc, "--matchmaking-url=URL    Tracker base URL, or a dead URL to disable it")
    add_bullet(doc, "--public-address=IP:port Pin Glyph + tracker address (localhost tests only)")
    add_bullet(doc, "--no-cache-connect       Do not dial known-peers.json on startup")
    add_bullet(doc, "--clear-peer-cache       Delete known-peers.json")
    add_bullet(doc, "--spawn-bots=N           Used by launch-full-test for Player 1")

    add_heading(doc, "11. Troubleshooting", 1)
    add_bullet(doc, "Window flashes and closes: the .ps1 failed to parse or the exe is missing. Run the .bat from a cmd window to see the error. Scripts must stay ASCII; Windows PowerShell 5.1 cannot read UTF-8 em-dashes.")
    add_bullet(doc, "Exe not found: run the matching build_all_*.bat in the same folder family (release vs develop).")
    add_bullet(doc, "Peers invisible: same-machine WAN Glyph/tracker addresses hairpin-fail. Use a localhost-pinned script.")
    add_bullet(doc, "They connect to the wrong person: clear cache (--clear-peer-cache) after an internet test.")
    add_bullet(doc, "Port in use: close leftover Carcosa.exe / Carcosa.Matchmaking.exe from a previous run.")

    add_heading(doc, "12. Confirmation checklist", 1)
    add_para(doc, "Walk each row, paste a screenshot in the placeholder above that script, then mark it confirmed.")
    add_bullet(doc, "[ ] release/build_all_release.bat")
    add_bullet(doc, "[x] release/launch-full-test.bat (stack launches)")
    add_bullet(doc, "[x] release/launch-two-players-local-tracker.bat")
    add_bullet(doc, "[ ] release/launch-two-players-local-tracker-no-cache.bat")
    add_bullet(doc, "[x] release/launch-two-players-no-tracker-no-cache.bat")
    add_bullet(doc, "[ ] release/launch-two-players.bat")
    add_bullet(doc, "[ ] release/launch-two-players-no-cache.bat")
    add_bullet(doc, "[ ] release/prepare-release.bat")
    add_bullet(doc, "[ ] develop/build_all_debug.bat")
    add_bullet(doc, "[ ] develop copies of each launcher")

    add_heading(doc, "13. Related docs", 1)
    add_bullet(doc, "README.md - architecture and current gameplay systems")
    add_bullet(doc, "docs/REMOTE_TESTING_GUIDE.md - earlier remote-peer notes (this Word doc supersedes the 'how to run scripts' part)")
    add_bullet(doc, "backlog/README.md - ticket-sized remaining work")
    add_bullet(doc, "implementations/ - plans (NAT/TURN, loot, overworld vision, mesh)")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
